from langchain_core.document_loaders import BaseBlobParser, Blob
from langchain_community.document_loaders.blob_loaders import FileSystemBlobLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from langgraph.graph import Graph, START, END
from openai import AzureOpenAI
import os
from pathlib import Path
from dotenv import load_dotenv

# 환경변수 로드
load_dotenv(Path(__file__).resolve().parent.parent.parent / ".env")

# Azure OpenAI 클라이언트 설정
client = AzureOpenAI(
    api_key=os.getenv('AZURE_OPENAI_API_KEY'),
    api_version="2024-02-01",  # 최신 API 버전
    azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT')  
)


# --------- Step 1: 파일 로드와 파싱 ---------

# 각 파일 유형별 파서 정의
class PDFParser(BaseBlobParser):
    def lazy_parse(self, blob: Blob):
        from langchain_community.document_loaders import PyMuPDFLoader
        loader = PyMuPDFLoader(blob.source)
        yield from loader.lazy_load()

class TXTParser(BaseBlobParser):
    def lazy_parse(self, blob: Blob):
        with blob.as_bytes_io() as f:
            yield Document(page_content=f.read().decode("utf-8"), metadata={"source": blob.source})

class DOCXParser(BaseBlobParser):
    def lazy_parse(self, blob: Blob):
        import docx
        doc = docx.Document(blob.source)
        for para in doc.paragraphs:
            yield Document(page_content=para.text, metadata={"source": blob.source})

class CSVParser(BaseBlobParser):
    def lazy_parse(self, blob: Blob):
        import pandas as pd
        df = pd.read_csv(blob.source)
        for _, row in df.iterrows():
            yield Document(page_content=str(row.to_dict()), metadata={"source": blob.source})

# 통합 BlobLoader 생성
blob_loader = FileSystemBlobLoader(path="C:/Users/Administrator/.vscode/Staff-Handover-Agent/data", glob="*.txt")
blobs = list(blob_loader.yield_blobs())

# 파일 확장자별 파서 매핑
parsers = {
    ".pdf": PDFParser(),
    ".txt": TXTParser(),
    ".docx": DOCXParser(),
    ".csv": CSVParser()
}

# 문서 파싱 및 통합
all_documents = []
for blob in blobs:
    ext = blob.source.split('.')[-1].lower()
    parser = parsers.get(f".{ext}")
    if parser:
        all_documents.extend(parser.lazy_parse(blob))
    else:
        print(f"Unsupported file format: {ext}")

# --------- Step 2: 문서 분할 ---------
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
split_documents = text_splitter.split_documents(all_documents)

# --------- Step 3: VectorStore 및 검색기 생성 ---------
embeddings = AzureOpenAIEmbeddings(
    model="text-embedding-3-large", # 모델 이름
    openai_api_version="2024-02-01",
    api_key="YOUR_AZURE_API_KEY",
    azure_endpoint="https://your-endpoint.openai.azure.com/"
)
vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)
retriever = vectorstore.as_retriever()

# --------- Step 4: 일정 추출 LLM 체인 ---------

# 일정 데이터 모델 정의
class Schedule(BaseModel):
    title: str = Field(..., description="일정 제목")
    date: str = Field(..., description="날짜 (YYYY-MM-DD)")
    time: str = Field(..., description="시간 (HH:MM)")

parser = PydanticOutputParser(pydantic_object=Schedule)

prompt = ChatPromptTemplate.from_messages([
    ("system", "아래 텍스트에서 일정을 추출하세요:\n{format_instructions}"),
    ("human", "{text}")
]).partial(format_instructions=parser.get_format_instructions())

llm = AzureChatOpenAI(
    azure_deployment="your-deployment",
    openai_api_version="2024-02-01",
    temperature=0.0,
    api_key="YOUR_AZURE_API_KEY",
    azure_endpoint="https://your-endpoint.openai.azure.com/"
)

schedule_chain = prompt | llm | parser

# --------- Step 5: 일정 데이터 처리 ---------
schedules = []
for doc in split_documents:
    result = schedule_chain.invoke({"text": doc.page_content})
    schedules.append(result)

# --------- Step 6: 일정 시각화 ---------

def create_schedule_graph(data):
    graph = Graph()
    graph.add_node(START, label="시작")
    for idx, schedule in enumerate(data):
        node_name = f"schedule_{idx}"
        label = f"{schedule.title}\n{schedule.date} {schedule.time}"
        graph.add_node(node_name, label=label)
        graph.add_edge(START if idx == 0 else f"schedule_{idx-1}", node_name)
    graph.add_node(END, label="끝")
    graph.add_edge(f"schedule_{len(data)-1}", END)
    return graph

schedule_graph = create_schedule_graph(schedules)
schedule_graph.get("graph").draw_mermaid_png("schedule_graph.png")
